import streamlit as st
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import base64
import io
from PyPDF2 import PdfReader
import docx
import re
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from typing import Tuple, Optional, Dict, List, Any
import time
import logging
from contextlib import contextmanager

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# App configuration
st.set_page_config(
    page_title="ExamPredictor AI",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Constants
MODEL_NAME = "microsoft/bitnet-b1.58-2B-4T"
TOPICS = ["Quadratic Equations", "Linear Equations", "Geometry", "Trigonometry", 
          "Calculus", "Statistics", "Probability", "Matrices"]
MAX_CONTENT_LENGTH = 4000  # Slightly increased limit
MAX_QUESTIONS = 15  # Increased max questions
DEFAULT_QUESTIONS = 5  # Increased default
DIFFICULTY_LEVELS = ["Basic", "Intermediate", "Advanced", "Mixed"]
EXAM_TYPES = ["Final Exam", "Midterm", "Quiz", "Assessment", "Practice Test", "Homework"]

# CSS styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1.5rem;
    }
    .stButton button {
        width: 100%;
    }
    .subheader {
        color: #0D47A1;
        border-bottom: 1px solid #E0E0E0;
        padding-bottom: 0.5rem;
    }
    .info-box {
        background-color: #F3F6F9;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1E88E5;
    }
    .question-box {
        background-color: #F9F9F9;
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 0.8rem;
        border-left: 3px solid #4CAF50;
    }
</style>
""", unsafe_allow_html=True)

# Progress indicator context manager
@contextmanager
def progress_indicator(message: str):
    """Context manager to show progress with spinner and timing."""
    start_time = time.time()
    spinner_placeholder = st.empty()
    spinner_placeholder.info(f"{message}...")
    try:
        yield
    finally:
        end_time = time.time()
        duration = end_time - start_time
        spinner_placeholder.success(f"{message} completed in {duration:.2f} seconds")

# Cache model loading
@st.cache_resource(show_spinner=False)
def load_model() -> Tuple[Optional[AutoTokenizer], Optional[AutoModelForCausalLM], Optional[str]]:
    """Load BitNet b1.58-2B-4T with quantization for efficiency."""
    try:
        with progress_indicator("Loading AI model"):
            tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, trust_remote_code=True)
            model = AutoModelForCausalLM.from_pretrained(
                MODEL_NAME,
                device_map="auto",
                torch_dtype=torch.float16,
                load_in_4bit=True,  # 4-bit quantization via bitsandbytes
                trust_remote_code=True
            )
            logger.info(f"Model {MODEL_NAME} loaded successfully")
            return tokenizer, model, None
    except Exception as e:
        logger.error(f"Failed to load model: {str(e)}")
        return None, None, f"Failed to load model: {str(e)}"

# File processing
def extract_text_from_pdf(file: io.BytesIO) -> str:
    """Extract text from a PDF file."""
    try:
        pdf_reader = PdfReader(file)
        text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
        logger.info(f"PDF extraction: {len(text)} characters from {len(pdf_reader.pages)} pages")
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return ""

def extract_text_from_docx(file: io.BytesIO) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file)
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        logger.info(f"DOCX extraction: {len(text)} characters from {len(doc.paragraphs)} paragraphs")
        return text
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return ""

def extract_text_from_file(file: st.uploaded_file_manager.UploadedFile) -> Tuple[str, Optional[str]]:
    """Extract text from uploaded file based on type."""
    try:
        file_type = file.type
        file_content = io.BytesIO(file.read())
        
        if file_type == "application/pdf":
            return extract_text_from_pdf(file_content), None
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_text_from_docx(file_content), None
        elif file_type == "text/plain":
            return file_content.read().decode("utf-8"), None
        else:
            error_msg = f"Unsupported file type: {file_type}"
            logger.warning(error_msg)
            return "", error_msg
    except Exception as e:
        error_msg = f"Error processing file: {str(e)}"
        logger.error(error_msg)
        return "", error_msg

def clean_and_format_text(text: str) -> str:
    """Clean and format extracted text."""
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove unnecessary spaces
    text = re.sub(r' {2,}', ' ', text)
    # Limit length
    if len(text) > MAX_CONTENT_LENGTH:
        text = text[:MAX_CONTENT_LENGTH]
        logger.info(f"Text truncated to {MAX_CONTENT_LENGTH} characters")
    
    return text.strip()

# Question generation
def generate_questions(
    tokenizer: AutoTokenizer,
    model: AutoModelForCausalLM,
    course_info: str,
    exam_type: str,
    difficulty: str,
    num_questions: int,
    topic: str,
    content: str,
    temperature: float,
    max_tokens: int,
    include_answers: bool
) -> Tuple[str, Optional[str]]:
    """Generate exam questions using BitNet."""
    try:
        # Clean content
        content = clean_and_format_text(content)
        
        # Create prompt with improved instructions
        answer_instruction = "Include the correct answer for each question." if include_answers else ""
        difficulty_instruction = (
            "Mix questions of varying difficulty levels." 
            if difficulty == "Mixed" 
            else f"All questions should be {difficulty.lower()} level."
        )
        
        prompt = f"""You are an expert professor in {course_info or 'the subject'}.

Create {num_questions} high-quality questions for a {exam_type.lower()}.
{difficulty_instruction}
{"Focus specifically on " + topic + "." if topic != 'General' else 'Cover various topics from the provided course material.'}
{answer_instruction}

{"Course Material: " + content if content else ""}

Format your response as:
1. First question
2. Second question
...

{"If providing answers, add them after each question with 'Answer:' prefix." if include_answers else ""}
"""
        # Tokenize with dynamic handling for long prompts
        try:
            max_input_length = 1024 if len(content) > 2000 else 512
            inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=max_input_length).to(model.device)
            
            # Generate with improved parameters
            with torch.no_grad():
                outputs = model.generate(
                    **inputs,
                    max_new_tokens=max_tokens,
                    temperature=temperature,
                    top_p=0.92,  # Slightly adjusted for better quality
                    top_k=50,    # Added parameter for more focused outputs
                    do_sample=True,
                    repetition_penalty=1.1  # Reduce repetitive responses
                )
            
            # Decode response
            generated_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
            response = generated_text[len(tokenizer.decode(inputs.input_ids[0], skip_special_tokens=True)):].strip()
            
            # Fix numbering if needed
            if re.match(r'^\s*\d+\.', response):
                response = re.sub(r'^\s*\d+\.', '1.', response, count=1)
                
            logger.info(f"Generated {len(response)} characters of questions")
            return response, None
            
        except torch.cuda.OutOfMemoryError:
            return "", "GPU memory exceeded. Try reducing content length or generating fewer questions."
        
    except Exception as e:
        logger.error(f"Generation error: {str(e)}")
        return "", f"Error generating questions: {str(e)}"

# Topic analysis
def analyze_topics(text: str) -> pd.DataFrame:
    """Analyze text to estimate topic probabilities."""
    # Define keyword mapping for rudimentary topic analysis
    topic_keywords = {
        "Quadratic Equations": ["quadratic", "x²", "x^2", "parabola", "discriminant", "roots", "completing the square"],
        "Linear Equations": ["linear", "slope", "y-intercept", "solve for x", "direct variation", "system of equations"],
        "Geometry": ["area", "volume", "perimeter", "triangle", "circle", "polygon", "angle", "diameter"],
        "Trigonometry": ["sin", "cos", "tan", "angle", "radian", "degree", "pythagorean"],
        "Calculus": ["derivative", "integral", "limit", "differentiate", "integrate", "rate of change"],
        "Statistics": ["mean", "median", "mode", "standard deviation", "normal distribution", "probability"],
        "Probability": ["chance", "likelihood", "random", "sample space", "event", "outcome", "probability"],
        "Matrices": ["matrix", "determinant", "eigenvalue", "vector", "linear algebra", "inverse"]
    }
    
    # Count keyword occurrences
    counts = {topic: 0 for topic in topic_keywords}
    text_lower = text.lower()
    
    for topic, keywords in topic_keywords.items():
        for keyword in keywords:
            counts[topic] += text_lower.count(keyword.lower())
    
    # Calculate probabilities with smoothing
    total = sum(counts.values()) + 1e-6  # Avoid division by zero
    probs = {topic: count/total for topic, count in counts.items()}
    
    # If no significant matches, use uniform distribution
    if max(probs.values()) < 0.1:
        probs = {topic: 1.0/len(topic_keywords) for topic in topic_keywords}
    
    # Create DataFrame
    df = pd.DataFrame({
        "Topic": list(probs.keys()),
        "Probability": list(probs.values())
    })
    
    return df.sort_values(by="Probability", ascending=False)

def create_topic_chart(probabilities: pd.DataFrame) -> Any:
    """Create a horizontal bar chart for topic probabilities."""
    fig, ax = plt.subplots(figsize=(10, 5))
    
    # Sort by probability
    df_sorted = probabilities.sort_values(by="Probability")
    
    # Plot horizontal bars
    bars = ax.barh(df_sorted["Topic"], df_sorted["Probability"], color="#1E88E5")
    
    # Add values as text
    for i, bar in enumerate(bars):
        width = bar.get_width()
        ax.text(
            width + 0.01, 
            bar.get_y() + bar.get_height()/2, 
            f"{width:.2f}", 
            va='center'
        )
    
    # Style the chart
    ax.set_xlabel("Probability")
    ax.set_title("Topic Distribution Analysis")
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    return fig

# Format questions for display
def format_questions_for_display(questions_text: str) -> str:
    """Format questions as HTML for better display."""
    if not questions_text:
        return ""
    
    # Split by question number pattern
    questions = re.split(r'(?:\n|^)(\d+\.\s*)', questions_text)
    formatted_html = ""
    
    # Skip the first empty split if it exists
    start_idx = 0
    if questions and not questions[0].strip():
        start_idx = 1
    
    # Process pairs of (number, content)
    for i in range(start_idx, len(questions), 2):
        if i+1 < len(questions):
            number = questions[i]
            content = questions[i+1]
            
            # Split content and answer if exists
            question_parts = content.split("Answer:", 1)
            question_content = question_parts[0].strip()
            
            answer_html = ""
            if len(question_parts) > 1:
                answer_content = question_parts[1].strip()
                answer_html = f'<div style="margin-top:0.5rem;padding-left:1rem;border-left:2px solid #FFA000;"><strong>Answer:</strong> {answer_content}</div>'
            
            # Format as HTML with question box styling
            formatted_html += f"""
            <div class="question-box">
                <strong>{number}</strong>{question_content}
                {answer_html}
            </div>
            """
    
    return formatted_html

# Save questions to CSV
def format_questions_for_csv(questions_text: str) -> str:
    """Format questions for CSV download."""
    lines = questions_text.split("\n")
    csv_rows = ["Question,Answer"]
    
    current_question = ""
    current_answer = ""
    
    for line in lines:
        if re.match(r'^\d+\.', line):
            # If we have a previous question, add it to our CSV rows
            if current_question:
                csv_rows.append(f"\"{current_question}\",\"{current_answer}\"")
                
            # Start a new question
            current_question = line.strip()
            current_answer = ""
        elif "Answer:" in line:
            # Extract answer
            current_answer = line.split("Answer:", 1)[1].strip()
        elif current_question:
            # Continue previous question or answer
            if current_answer:
                current_answer += " " + line.strip()
            else:
                current_question += " " + line.strip()
    
    # Add the last question
    if current_question:
        csv_rows.append(f"\"{current_question}\",\"{current_answer}\"")
    
    return "\n".join(csv_rows)

# Main UI
def main():
    """Render the Streamlit app."""
    # Header
    st.markdown("<h1 class='main-header'>📝 Exam Question Predictor</h1>", unsafe_allow_html=True)
    
    # Initialize session state
    if 'generated_questions' not in st.session_state:
        st.session_state.generated_questions = ""
    if 'topic_probs' not in st.session_state:
        st.session_state.topic_probs = None
    if 'content' not in st.session_state:
        st.session_state.content = ""
    if 'generation_time' not in st.session_state:
        st.session_state.generation_time = 0
    
    # Load model
    tokenizer, model, error = load_model()
    if error:
        st.error(error)
        st.error("Try refreshing the page or contact support.")
        return

    # Layout
    col1, col2 = st.columns([3, 2], gap="medium")

    # Input Section
    with col1:
        st.markdown("<h2 class='subheader'>Input Course Material</h2>", unsafe_allow_html=True)
        
        # Input method
        input_method = st.radio("Choose input method:", ["Upload File", "Enter Text"], horizontal=True)
        
        content = st.session_state.content
        if input_method == "Upload File":
            uploaded_file = st.file_uploader(
                "Upload syllabus, notes, or textbook",
                type=["pdf", "txt", "docx"],
                accept_multiple_files=False,
                help="Supported formats: PDF, DOCX, TXT"
            )
            if uploaded_file:
                with progress_indicator("Extracting text"):
                    content, error = extract_text_from_file(uploaded_file)
                    if error:
                        st.error(error)
                    elif content:
                        st.session_state.content = content
                        st.success(f"Successfully extracted {len(content)} characters")
                        
                        # Show text analysis automatically
                        if len(content) > 100:
                            st.session_state.topic_probs = analyze_topics(content)
                        
                        with st.expander("Preview Content", expanded=False):
                            st.text_area("Extracted text:", value=content[:2000] + ("..." if len(content) > 2000 else ""), height=200, disabled=True)
                    else:
                        st.warning("No text could be extracted from this file. Try another file or enter text manually.")
        else:
            content = st.text_area(
                "Enter course material (e.g., syllabus, notes):",
                height=250,
                value=st.session_state.content,
                placeholder="Example: The quadratic formula x = (-b ± √(b² - 4ac)) / 2a allows us to solve any quadratic equation in the form ax² + bx + c = 0..."
            )
            st.session_state.content = content
            
            # Analyze button for manual text
            if content and st.button("Analyze Text", type="secondary"):
                with progress_indicator("Analyzing content"):
                    st.session_state.topic_probs = analyze_topics(content)
                    st.success("Content analyzed! See topic distribution in the output section.")
        
        # Course and exam details
        st.markdown("<h2 class='subheader'>Exam Details</h2>", unsafe_allow_html=True)
        
        col_subject, col_exam = st.columns(2)
        with col_subject:
            course_info = st.text_input(
                "Course/Subject Name", 
                value="Algebra", 
                placeholder="e.g., Algebra, Physics, Calculus"
            )
        with col_exam:
            exam_type = st.selectbox("Exam Type", EXAM_TYPES)
        
        col_diff, col_quest = st.columns(2)
        with col_diff:
            difficulty = st.selectbox("Question Difficulty", DIFFICULTY_LEVELS)
        with col_quest:
            num_questions = st.number_input(
                "Number of Questions",
                min_value=1,
                max_value=MAX_QUESTIONS,
                value=DEFAULT_QUESTIONS,
                step=1
            )
        
        col_topic, col_ans = st.columns(2)
        with col_topic:
            # Auto-select topic based on analysis if available
            default_topic = "General"
            if st.session_state.topic_probs is not None and not st.session_state.topic_probs.empty:
                default_topic = st.session_state.topic_probs.iloc[0]["Topic"]
                
            topic = st.selectbox(
                "Focus Topic", 
                ["General"] + TOPICS,
                index=0 if default_topic == "General" else TOPICS.index(default_topic) + 1
            )
        with col_ans:
            include_answers = st.checkbox("Include Answers", value=True)
        
        # Advanced settings
        with st.expander("Advanced Settings", expanded=False):
            temperature = st.slider(
                "Creativity (Temperature)", 
                0.1, 1.0, 0.7, 0.1, 
                help="Higher values (>0.7) = more creative but potentially less accurate questions"
            )
            max_tokens = st.slider(
                "Max Response Length", 
                256, 2048, 768, 
                help="Controls output length. Use higher values for more detailed or numerous questions"
            )
    
    # Output Section
    with col2:
        st.markdown("<h2 class='subheader'>Generated Questions</h2>", unsafe_allow_html=True)
        
        # Topic probabilities if available
        if st.session_state.topic_probs is not None and not st.session_state.topic_probs.empty:
            with st.expander("Topic Distribution Analysis", expanded=True):
                st.dataframe(
                    st.session_state.topic_probs,
                    use_container_width=True,
                    hide_index=True
                )
                chart = create_topic_chart(st.session_state.topic_probs)
                st.pyplot(chart)
        
        # Generate button
        generate_disabled = not (content.strip() or topic != "General") or not tokenizer or not model
        generate_button = st.button(
            "Generate Questions", 
            type="primary", 
            disabled=generate_disabled,
            use_container_width=True
        )
        
        # Show info if button is disabled
        if generate_disabled:
            st.markdown(
                "<div class='info-box'>Please enter course material or select a specific topic to generate questions.</div>", 
                unsafe_allow_html=True
            )
        
        # Generate questions when button is clicked
        if generate_button:
            start_time = time.time()
            with st.spinner("Generating questions. This may take 10-30 seconds..."):
                response, error = generate_questions(
                    tokenizer, model, course_info, exam_type, difficulty,
                    num_questions, topic, content, temperature, max_tokens,
                    include_answers
                )
                
                end_time = time.time()
                st.session_state.generation_time = end_time - start_time
                st.session_state.generated_questions = response
                
                if error:
                    st.error(error)
        
        # Display questions if available
        if st.session_state.generated_questions:
            st.success(f"Questions generated in {st.session_state.generation_time:.1f} seconds")
            
            # Format and display questions
            formatted_questions = format_questions_for_display(st.session_state.generated_questions)
            st.markdown(formatted_questions, unsafe_allow_html=True)
            
            # Download options
            col_csv, col_txt = st.columns(2)
            with col_csv:
                csv_content = format_questions_for_csv(st.session_state.generated_questions)
                filename_csv = f"{course_info or 'exam'}_{topic.lower().replace(' ', '_')}_questions.csv"
                st.download_button(
                    label="Download as CSV",
                    data=csv_content,
                    file_name=filename_csv,
                    mime="text/csv",
                    type="secondary",
                    use_container_width=True
                )
            
            with col_txt:
                filename_txt = f"{course_info or 'exam'}_{topic.lower().replace(' ', '_')}_questions.txt"
                st.download_button(
                    label="Download as TXT",
                    data=st.session_state.generated_questions,
                    file_name=filename_txt,
                    mime="text/plain",
                    type="secondary",
                    use_container_width=True
                )

    # Sidebar
    with st.sidebar:
        st.markdown("<h2 class='subheader'>About ExamPredictor AI</h2>", unsafe_allow_html=True)
        st.markdown("""
        <div class="info-box">
        This tool generates exam questions based on course materials or selected topics using the BitNet b1.58-2B-4T AI model.
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("**Features:**")
        st.markdown("""
        - Upload PDF, TXT, or DOCX files
        - Generate questions with optional answers
        - Real topic analysis and distribution
        - Multiple download formats
        - Adjustable question count and difficulty
        """)
        
        st.markdown("<h2 class='subheader'>Tips for Best Results</h2>", unsafe_allow_html=True)
        st.markdown("""
        - Upload detailed notes or syllabus for better context.
        - Select a specific topic to focus questions.
        - Adjust difficulty to match your exam needs.
        - Use temperature >0.7 for creative questions.
        - Include answers to provide solution examples.
        - Try "Mixed" difficulty for a comprehensive range.
        """)
        
        st.markdown("<h2 class='subheader'>Models & References</h2>", unsafe_allow_html=True)
        st.markdown("""
        - Using Microsoft's BitNet b1.58-2B-4T model
        - Optimized with 4-bit quantization
        - Built with Streamlit and HuggingFace Transformers
        """)

if __name__ == "__main__":
    main()
